#!/usr/bin/env python3
"""
SCS Checker - Enhanced Export Module
Provides PDF, Excel, and Word document export for scan results.
"""
import json
import os
import io
from datetime import datetime


def export_excel(scan_data, project_name, scan_time):
    """Export scan results to Excel format using openpyxl."""
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        # Fallback: generate CSV
        return export_csv(scan_data, project_name, scan_time)

    wb = Workbook()

    # Style definitions
    header_font = Font(name='Calibri', size=12, bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='1F4E79', end_color='1F4E79', fill_type='solid')
    cell_font = Font(name='Calibri', size=10)
    border = Border(
        side=Side(style='thin', color='B0B0B0'),
        outline=Side(style='thin', color='B0B0B0')
    )
    sev_colors = {
        'critical': PatternFill(start_color='FF4444', end_color='FF4444', fill_type='solid'),
        'high': PatternFill(start_color='FF8888', end_color='FF8888', fill_type='solid'),
        'medium': PatternFill(start_color='FFCC44', end_color='FFCC44', fill_type='solid'),
        'low': PatternFill(start_color='88CC88', end_color='88CC88', fill_type='solid'),
    }

    # === Sheet 1: Summary ===
    ws1 = wb.active
    ws1.title = '检测摘要'

    ws1['A1'] = 'SCS Checker - 供应链安全检测报告'
    ws1['A1'].font = Font(name='Calibri', size=16, bold=True, color='1F4E79')
    ws1.merge_cells('A1:D1')

    ws1['A3'] = '项目名称'
    ws1['B3'] = project_name
    ws1['A4'] = '检测时间'
    ws1['B4'] = scan_time
    ws1['A5'] = '总包数'
    ws1['B5'] = scan_data.get('total_packages', 0)
    ws1['A6'] = '漏洞包数'
    ws1['B6'] = scan_data.get('vulnerable_packages', 0)
    ws1['A7'] = '漏洞总数'
    ws1['B7'] = scan_data.get('total_vulnerabilities', 0)

    sev = scan_data.get('severity_counts', {})
    ws1['A9'] = '严重性分布'
    ws1['A9'].font = Font(bold=True, size=12)
    ws1['A10'] = 'Critical'
    ws1['B10'] = sev.get('critical', 0)
    ws1['A11'] = 'High'
    ws1['B11'] = sev.get('high', 0)
    ws1['A12'] = 'Medium'
    ws1['B12'] = sev.get('medium', 0)
    ws1['A13'] = 'Low'
    ws1['B13'] = sev.get('low', 0)

    for row in range(3, 14):
        ws1[f'A{row}'].font = Font(bold=True)

    ws1.column_dimensions['A'].width = 20
    ws1.column_dimensions['B'].width = 40

    # === Sheet 2: Vulnerability Details ===
    ws2 = wb.create_sheet('漏洞详情')
    headers = ['包名', '当前版本', '漏洞ID', '严重性', 'CVSS评分', '摘要', '修复版本']
    for col, h in enumerate(headers, 1):
        cell = ws2.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = border

    row = 2
    for pkg in scan_data.get('packages', []):
        if pkg.get('vuln_count', 0) == 0:
            continue
        fix_ver = ''
        for vuln in pkg.get('vulnerabilities', []):
            affected = vuln.get('affected_versions', {})
            fixes = affected.get('fix_versions', []) if affected else []
            if fixes:
                fix_ver = ', '.join(fixes)

            ws2.cell(row=row, column=1, value=pkg['package']).font = cell_font
            ws2.cell(row=row, column=2, value=pkg.get('version', '')).font = cell_font
            ws2.cell(row=row, column=3, value=vuln.get('id', '')).font = cell_font
            sev_cell = ws2.cell(row=row, column=4, value=vuln.get('severity', 'unknown'))
            sev_cell.font = Font(bold=True, color='FFFFFF')
            sev_fill = sev_colors.get(vuln.get('severity', ''), PatternFill())
            sev_cell.fill = sev_fill
            sev_cell.alignment = Alignment(horizontal='center')
            ws2.cell(row=row, column=5, value=vuln.get('cvss_score', 0)).font = cell_font
            ws2.cell(row=row, column=6, value=vuln.get('summary', '')[:200]).font = cell_font
            ws2.cell(row=row, column=7, value=fix_ver).font = cell_font

            for col in range(1, 8):
                ws2.cell(row=row, column=col).border = border

            row += 1

    # Auto-width
    for col in range(1, 8):
        ws2.column_dimensions[get_column_letter(col)].width = [20, 15, 25, 12, 12, 60, 20][col-1]

    # === Sheet 3: All Packages ===
    ws3 = wb.create_sheet('所有包')
    headers3 = ['包名', '版本', '是否直接依赖', '漏洞数', '生态系统']
    for col, h in enumerate(headers3, 1):
        cell = ws3.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal='center')

    for i, pkg in enumerate(scan_data.get('packages', []), 2):
        ws3.cell(row=i, column=1, value=pkg.get('package', ''))
        ws3.cell(row=i, column=2, value=pkg.get('version', ''))
        ws3.cell(row=i, column=3, value='是' if pkg.get('is_direct') else '否')
        ws3.cell(row=i, column=4, value=pkg.get('vuln_count', 0))
        ws3.cell(row=i, column=5, value=pkg.get('ecosystem', 'PyPI'))

    for col in range(1, 6):
        ws3.column_dimensions[get_column_letter(col)].width = [25, 15, 15, 10, 12][col-1]

    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return buf


def export_csv(scan_data, project_name, scan_time):
    """Fallback CSV export."""
    output = []
    output.append(f'# SCS Checker Report - {project_name} - {scan_time}')
    output.append(f'# Total Packages: {scan_data.get("total_packages", 0)}')
    output.append(f'# Vulnerable Packages: {scan_data.get("vulnerable_packages", 0)}')
    output.append('')
    output.append('Package,Version,VulnID,Severity,CVSS,Summary,FixVersion')

    for pkg in scan_data.get('packages', []):
        if pkg.get('vuln_count', 0) == 0:
            continue
        for vuln in pkg.get('vulnerabilities', []):
            fix_ver = ''
            affected = vuln.get('affected_versions', {})
            fixes = affected.get('fix_versions', []) if affected else []
            if fixes:
                fix_ver = ', '.join(fixes)
            summary = vuln.get('summary', '').replace(',', ';')[:200]
            output.append(f'{pkg["package"]},{pkg.get("version","")},{vuln.get("id","")},{vuln.get("severity","")},{vuln.get("cvss_score",0)},{summary},{fix_ver}')

    buf = io.BytesIO('\n'.join(output).encode('utf-8-sig'))
    return buf


def export_word(scan_data, project_name, scan_time, risk_score=0):
    """Export scan results to Word .docx format."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return None

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading('SCS Checker 供应链安全检测报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    doc.add_paragraph('')
    meta = doc.add_table(rows=4, cols=2)
    meta.style = 'Light Grid Accent 1'
    meta_data = [
        ('项目名称', project_name),
        ('检测时间', scan_time),
        ('风险评分', str(risk_score)),
        ('报告生成', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
    ]
    for i, (k, v) in enumerate(meta_data):
        meta.cell(i, 0).text = k
        meta.cell(i, 1).text = str(v)

    doc.add_paragraph('')

    # Summary
    doc.add_heading('一、检测摘要', level=1)
    sev = scan_data.get('severity_counts', {})
    summary = doc.add_table(rows=7, cols=2)
    summary.style = 'Light Grid Accent 1'
    summary_data = [
        ('总包数', str(scan_data.get('total_packages', 0))),
        ('漏洞包数', str(scan_data.get('vulnerable_packages', 0))),
        ('漏洞总数', str(scan_data.get('total_vulnerabilities', 0))),
        ('Critical', str(sev.get('critical', 0))),
        ('High', str(sev.get('high', 0))),
        ('Medium', str(sev.get('medium', 0))),
        ('Low', str(sev.get('low', 0))),
    ]
    for i, (k, v) in enumerate(summary_data):
        summary.cell(i, 0).text = k
        summary.cell(i, 1).text = v

    # Vulnerability details
    doc.add_heading('二、漏洞详情', level=1)
    vuln_packages = [p for p in scan_data.get('packages', []) if p.get('vuln_count', 0) > 0]

    for pkg in vuln_packages:
        doc.add_heading(f'{pkg["package"]} ({pkg.get("version", "")})', level=2)
        for vuln in pkg.get('vulnerabilities', []):
            p = doc.add_paragraph()
            p.add_run(f'  {vuln.get("id", "")} ').bold = True
            sev = vuln.get('severity', 'unknown')
            run = p.add_run(f'[{sev.upper()}] ')
            sev_colors = {
                'critical': RGBColor(0xDC, 0x26, 0x26),
                'high': RGBColor(0xEF, 0x44, 0x44),
                'medium': RGBColor(0xF5, 0x9E, 0x0B),
                'low': RGBColor(0x10, 0xB9, 0x81),
            }
            run.font.color.rgb = sev_colors.get(sev, RGBColor(0x6B, 0x72, 0x80))
            p.add_run(f'CVSS: {vuln.get("cvss_score", 0)} | ')

            affected = vuln.get('affected_versions', {})
            fixes = affected.get('fix_versions', []) if affected else []
            if fixes:
                p.add_run(f'修复版本: {", ".join(fixes)}').bold = True

            summary_text = vuln.get('summary', '')
            if summary_text:
                doc.add_paragraph(f'  摘要: {summary_text[:300]}', style='Intense Quote')

    # All packages
    doc.add_heading('三、依赖包列表', level=1)
    pkg_table = doc.add_table(rows=1, cols=4)
    pkg_table.style = 'Light Grid Accent 1'
    hdr = pkg_table.rows[0].cells
    hdr[0].text = '包名'
    hdr[1].text = '版本'
    hdr[2].text = '漏洞数'
    hdr[3].text = '直接依赖'

    for pkg in scan_data.get('packages', []):
        row = pkg_table.add_row().cells
        row[0].text = pkg.get('package', '')
        row[1].text = pkg.get('version', '')
        row[2].text = str(pkg.get('vuln_count', 0))
        row[3].text = '是' if pkg.get('is_direct') else '否'

    # Footer
    doc.add_paragraph('')
    doc.add_paragraph('— 本报告由 SCS Checker 自动生成 —').alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def export_pdf_html(scan_data, project_name, scan_time, risk_score=0):
    """Generate a print-friendly HTML page that can be saved as PDF.
    This avoids requiring wkhtmltopdf or reportlab on the server.
    """
    sev = scan_data.get('severity_counts', {})
    packages = scan_data.get('packages', [])
    vuln_packages = [p for p in packages if p.get('vuln_count', 0) > 0]

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<title>SCS Checker - {project_name} 安全报告</title>
<style>
@page {{ margin: 2cm; }}
body {{ font-family: 'Microsoft YaHei', 'SimSun', Arial, sans-serif; color: #333; line-height: 1.6; }}
h1 {{ text-align: center; color: #1F4E79; border-bottom: 3px solid #1F4E79; padding-bottom: 10px; }}
h2 {{ color: #1F4E79; margin-top: 20px; }}
table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
th {{ background: #1F4E79; color: #fff; padding: 8px; text-align: left; }}
td {{ border: 1px solid #ddd; padding: 6px 8px; }}
tr:nth-child(even) {{ background: #f5f5f5; }}
.sev-critical {{ color: #DC2626; font-weight: bold; }}
.sev-high {{ color: #EF4444; font-weight: bold; }}
.sev-medium {{ color: #F59E0B; font-weight: bold; }}
.sev-low {{ color: #10B981; font-weight: bold; }}
.meta {{ background: #f0f4f8; padding: 15px; border-radius: 8px; margin: 15px 0; }}
.watermark {{ position: fixed; top: 50%; left: 50%; transform: translate(-50%, -50%) rotate(-45deg);
  font-size: 80px; color: rgba(31, 78, 121, 0.05); z-index: -1; pointer-events: none; }}
@media print {{
  .no-print {{ display: none; }}
  body {{ font-size: 12px; }}
}}
</style>
</head>
<body>
<div class="watermark">SCS CHECKER</div>
<div class="no-print" style="text-align:center; margin-bottom:20px;">
  <button onclick="window.print()" style="padding:10px 30px; font-size:16px; background:#1F4E79; color:#fff; border:none; border-radius:8px; cursor:pointer;">
    打印 / 保存为PDF
  </button>
</div>

<h1>SCS Checker 供应链安全检测报告</h1>

<div class="meta">
  <strong>项目名称:</strong> {project_name}<br>
  <strong>检测时间:</strong> {scan_time}<br>
  <strong>风险评分:</strong> {risk_score}/100<br>
  <strong>报告生成:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
</div>

<h2>一、检测摘要</h2>
<table>
<tr><th>指标</th><th>数值</th></tr>
<tr><td>总包数</td><td>{scan_data.get('total_packages', 0)}</td></tr>
<tr><td>漏洞包数</td><td>{scan_data.get('vulnerable_packages', 0)}</td></tr>
<tr><td>漏洞总数</td><td>{scan_data.get('total_vulnerabilities', 0)}</td></tr>
<tr><td>Critical</td><td class="sev-critical">{sev.get('critical', 0)}</td></tr>
<tr><td>High</td><td class="sev-high">{sev.get('high', 0)}</td></tr>
<tr><td>Medium</td><td class="sev-medium">{sev.get('medium', 0)}</td></tr>
<tr><td>Low</td><td class="sev-low">{sev.get('low', 0)}</td></tr>
</table>

<h2>二、漏洞详情</h2>
'''

    for pkg in vuln_packages:
        html += f'<h3>{pkg["package"]} ({pkg.get("version", "")})</h3>\n<table>\n<tr><th>漏洞ID</th><th>严重性</th><th>CVSS</th><th>摘要</th><th>修复版本</th></tr>\n'
        for vuln in pkg.get('vulnerabilities', []):
            affected = vuln.get('affected_versions', {})
            fixes = affected.get('fix_versions', []) if affected else []
            fix_str = ', '.join(fixes) if fixes else '-'
            sev = vuln.get('severity', 'unknown')
            html += f'<tr><td>{vuln.get("id", "")}</td><td class="sev-{sev}">{sev.upper()}</td><td>{vuln.get("cvss_score", 0)}</td><td>{vuln.get("summary", "")[:200]}</td><td>{fix_str}</td></tr>\n'
        html += '</table>\n'

    html += '''
<h2>三、依赖包列表</h2>
<table>
<tr><th>包名</th><th>版本</th><th>漏洞数</th><th>直接依赖</th></tr>
'''
    for pkg in packages:
        html += f'<tr><td>{pkg.get("package", "")}</td><td>{pkg.get("version", "")}</td><td>{pkg.get("vuln_count", 0)}</td><td>{"是" if pkg.get("is_direct") else "否"}</td></tr>\n'

    html += '''</table>
<p style="text-align:center; color:#999; margin-top:30px;">— 本报告由 SCS Checker 自动生成 —</p>
</body>
</html>'''

    return html
